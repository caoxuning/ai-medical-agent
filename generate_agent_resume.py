"""
生成全新 Agent 方向简历
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='Microsoft YaHei', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_agent_resume():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # ========== 头部信息 ==========
    name = doc.add_paragraph()
    name_run = name.add_run('（姓名）')
    set_chinese_font(name_run, font_size=18, bold=True)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 求职意向
    target = doc.add_paragraph()
    target_run = target.add_run('求职意向：AI Agent 工程师 / 大模型应用开发')
    set_chinese_font(target_run, font_size=12, bold=True)
    target.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 联系信息
    contact = doc.add_paragraph()
    contact_run = contact.add_run('电话：（手机号） | 邮箱：（邮箱） |  GitHub：（GitHub链接）')
    set_chinese_font(contact_run, font_size=10)
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # ========== 教育背景 ==========
    heading = doc.add_paragraph()
    heading_run = heading.add_run('教育背景')
    set_chinese_font(heading_run, font_size=14, bold=True)
    
    edu = doc.add_paragraph()
    edu_text = '哈尔滨工业大学（威海） | 软件工程 | 本科 | 2023.09 - 2027.06（预计）'
    edu_run = edu.add_run(edu_text)
    set_chinese_font(edu_run)
    
    doc.add_paragraph()
    
    # ========== 专业技能 ==========
    heading = doc.add_paragraph()
    heading_run = heading.add_run('专业技能')
    set_chinese_font(heading_run, font_size=14, bold=True)
    
    skills = [
        ('Agent 架构', '熟练掌握 ReAct、Plan-and-Execute 等 Agent 设计模式，具备多轮对话状态管理、工具调用、记忆系统设计经验'),
        ('RAG 技术', '精通检索增强生成技术，包括向量检索（ChromaDB）、混合检索（BM25+向量）、RRF 融合、重排序优化'),
        ('大模型应用', '熟悉 GPT-4/GLM-4/Kimi 等 LLM API 调用，具备 Prompt Engineering、Function Calling、多 Agent 协作开发经验'),
        ('工程部署', '熟练使用 FastAPI/Flask 构建 API 服务，具备 Docker 容器化、Linux 服务器部署经验'),
        ('Python 技术栈', '熟练使用 PyTorch、Transformers、LangChain、NumPy、Pandas 等深度学习与数据处理库'),
    ]
    
    for skill_name, skill_desc in skills:
        p = doc.add_paragraph(style='List Bullet')
        name_run = p.add_run(f'{skill_name}：')
        set_chinese_font(name_run, bold=True)
        desc_run = p.add_run(skill_desc)
        set_chinese_font(desc_run)
    
    doc.add_paragraph()
    
    # ========== 项目经历 ==========
    heading = doc.add_paragraph()
    heading_run = heading.add_run('项目经历')
    set_chinese_font(heading_run, font_size=14, bold=True)
    
    # 项目 1
    p = doc.add_paragraph()
    title_run = p.add_run('企业级 RAG 合规文档智能问答系统')
    set_chinese_font(title_run, font_size=12, bold=True)
    
    p = doc.add_paragraph()
    info_run = p.add_run('独立项目 | 技术栈：FastAPI + ChromaDB + 智谱 AI + BM25 + JWT | GitHub：（链接）')
    set_chinese_font(info_run, font_size=9)
    
    proj1_points = [
        '设计并实现基于 RAG 架构的企业级文档问答系统，支持 PDF/Word/TXT/Markdown 多格式文档的智能解析与向量化，构建 136+ 文本块的企业合规知识库',
        '实现混合检索架构：向量检索（ChromaDB + 智谱 Embedding-3）+ 关键词检索（BM25）+ RRF 融合 + 轻量级重排序，实测 Recall@5 = 100%，MRR = 0.9667',
        '设计三级权限体系（Super Admin / Company Manager / Employee），基于 JWT 实现用户认证与资源隔离，支持邀请码机制的企业成员管理',
        '开发完整的 Web UI 与 RESTful API，实现文档上传、预览、问答交互，系统可直接交付企业使用',
    ]
    
    for point in proj1_points:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(point)
        set_chinese_font(run)
    
    doc.add_paragraph()
    
    # 项目 2
    p = doc.add_paragraph()
    title_run = p.add_run('基于 ReAct Agent 架构的智能医疗问诊系统')
    set_chinese_font(title_run, font_size=12, bold=True)
    
    p = doc.add_paragraph()
    info_run = p.add_run('独立项目 | 技术栈：ReAct + 智谱 GLM-4-Flash + 80万+ 医学知识库 + SQLite | GitHub：（链接）')
    set_chinese_font(info_run, font_size=9)
    
    proj2_points = [
        '设计并实现医疗领域 ReAct 风格 Agent，基于 Observation → Thought → Action 循环实现多轮问诊决策，接入智谱 GLM-4-Flash 作为推理引擎',
        '构建 80万+ 条医学知识库（覆盖内科、外科、儿科等 6 个科室），实现基于关键词的检索增强生成，为 Agent 提供专业医学知识支撑',
        '实现智能信息收集：自动提取症状、持续时间、疼痛位置/性质等信息，基于信息完整度判断（症状+持续时间+位置/性质）决定是否给出建议',
        '开发 Flask RESTful API 与响应式 Web 前端，支持会话管理与 SQLite 持久化存储，对话历史可跨会话恢复',
    ]
    
    for point in proj2_points:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(point)
        set_chinese_font(run)
    
    doc.add_paragraph()
    
    # ========== 技术亮点 ==========
    heading = doc.add_paragraph()
    heading_run = heading.add_run('技术亮点')
    set_chinese_font(heading_run, font_size=14, bold=True)
    
    highlights = [
        '独立完成 2 个完整的 Agent 项目，从架构设计到工程部署的全链路开发能力',
        '精通 RAG 技术栈，具备向量检索、混合检索、重排序等核心技术的实战经验',
        '熟悉大模型应用开发，具备 Prompt Engineering、Function Calling、多轮对话管理能力',
        '具备数据处理能力，完成 80万+ 条医学数据的清洗、解析与知识库构建',
    ]
    
    for highlight in highlights:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(highlight)
        set_chinese_font(run)
    
    doc.add_paragraph()
    
    # ========== 自我评价 ==========
    heading = doc.add_paragraph()
    heading_run = heading.add_run('自我评价')
    set_chinese_font(heading_run, font_size=14, bold=True)
    
    summary = doc.add_paragraph()
    summary_text = ('对 AI Agent 技术有浓厚兴趣，具备独立思考和解决问题的能力。'
                   '在 2 个独立项目中完成了从需求分析、技术选型、开发实现到部署交付的完整流程，'
                   '具备端到端的工程化落地能力。善于学习新技术，能够快速将前沿技术应用于实际项目。')
    summary_run = summary.add_run(summary_text)
    set_chinese_font(summary_run)
    
    # 保存
    output_path = 'resume_agent_template.docx'
    doc.save(output_path)
    print(f'✅ Agent 简历模板已生成: {output_path}')
    
    # 复制到桌面
    import shutil
    import os
    desktop_path = os.path.expanduser('~/Desktop/简历_Agent方向模板.docx')
    shutil.copy(output_path, desktop_path)
    print(f'✅ 已复制到桌面: {desktop_path}')

if __name__ == '__main__':
    create_agent_resume()

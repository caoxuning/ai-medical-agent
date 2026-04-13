"""
生成更新后的简历（Word格式）
添加两个Agent项目经历
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_resume():
    doc = Document()
    
    # 标题
    title = doc.add_heading('个人简历', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 基本信息
    doc.add_heading('基本信息', 1)
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Light Grid Accent 1'
    
    basic_info = [
        ['姓名：', '（请填写）', '出生年月：', '（请填写）'],
        ['民族：', '（请填写）', '身高：', '（请填写）'],
        ['电话：', '（请填写）', '政治面貌：', '（请填写）'],
        ['邮箱：', '（请填写）', '所在院校：', '哈尔滨工业大学'],
    ]
    
    for i, row_data in enumerate(basic_info):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
    
    doc.add_paragraph()
    
    # 教育背景
    doc.add_heading('教育背景', 1)
    edu = doc.add_paragraph()
    edu.add_run('2023.09 - 2027.06\n').bold = True
    edu.add_run('哈尔滨工业大学（威海）\n').bold = True
    edu.add_run('软件工程')
    
    # 专业技能
    doc.add_heading('专业技能', 1)
    skills = doc.add_paragraph(
        '算法框架：熟悉 PyTorch 构建深度学习模型，LSTM/注意力机制、GPT2/BERT 等 NLP 与时序模型；\n'
        '工程能力：熟悉数据预处理流水线搭建、模型训练监控（MAE/RMSE 等指标）、Flask API 部署；\n'
        'Agent/RAG：熟悉 RAG 检索增强架构，具备向量检索、混合检索（BM25+向量）、重排序优化实战经验；\n'
        '优化手段：熟悉 AdamW、学习率预热、正则化早停等策略，可独立解决过拟合、模型精度优化问题。'
    )
    
    # 项目经历
    doc.add_heading('项目经历', 1)
    
    # 项目1：企业级RAG合规文档智能问答系统
    doc.add_heading('企业级 RAG 合规文档智能问答系统', 2)
    p1 = doc.add_paragraph()
    p1.add_run('AI Agent 开发（独立项目）\n').bold = True
    
    p1.add_run('核心职责：').bold = True
    p1.add_run('独立设计并实现基于 RAG 架构的企业级文档问答系统，支持多格式文档智能解析、混合检索与权限隔离，为企业提供精准的合规知识服务。\n\n')
    
    responsibilities1 = [
        '构建 RAG 检索增强架构：向量检索（ChromaDB）+ 关键词检索（BM25）+ RRF 融合，引入轻量级重排序器，实测 Recall@5 = 100%，MRR = 0.9667；',
        '实现多格式文档处理：支持 PDF、Word、TXT、Markdown 多格式文档自动解析与向量化，实现文档分块、嵌入、存储全流程自动化；',
        '设计企业级权限管理：基于 FastAPI + JWT 实现 Super Admin / Company Manager / Employee 三级权限体系，支持邀请码机制的企业成员管理；',
        '完成工程化部署：开发 RESTful API 与完整前端界面，支持文档上传、预览、问答交互，输出标准化部署文档，系统可直接交付企业使用。'
    ]
    
    for resp in responsibilities1:
        doc.add_paragraph(resp, style='List Number')
    
    doc.add_paragraph()
    
    # 项目2：基于ReAct Agent架构的智能医疗问诊系统
    doc.add_heading('基于 ReAct Agent 架构的智能医疗问诊系统', 2)
    p2 = doc.add_paragraph()
    p2.add_run('AI Agent 开发（独立项目）\n').bold = True
    
    p2.add_run('核心职责：').bold = True
    p2.add_run('独立设计并实现医疗领域 ReAct 风格 Agent，集成智谱 AI 与 80万+ 医学知识库，实现症状收集、知识检索与建议生成的完整闭环。\n\n')
    
    responsibilities2 = [
        '构建 ReAct Agent 核心引擎：设计 Observation → Thought → Action 推理循环，Agent 基于对话上下文自主决策，接入智谱 GLM-4-Flash 作为推理引擎；',
        '构建医学知识库 RAG 系统：收集处理 80万+ 条真实医疗问答数据（覆盖内科、外科、儿科等6个科室），实现基于关键词的检索增强生成；',
        '实现智能信息收集与建议生成：设计症状提取、槽位填充机制，自动判断信息充足度（症状+持续时间+位置/性质），基于知识库返回专业医学建议；',
        '完成工程化部署：开发 Flask RESTful API，支持会话管理与 SQLite 持久化存储，实现可交互的医疗问诊服务，对话历史可跨会话恢复。'
    ]
    
    for resp in responsibilities2:
        doc.add_paragraph(resp, style='List Number')
    
    doc.add_paragraph()
    
    # 项目3：深度学习气象智能预测项目（原有）
    doc.add_heading('深度学习气象智能预测项目', 2)
    p3 = doc.add_paragraph()
    p3.add_run('AI 算法 Agent（负责人）\n').bold = True
    
    responsibilities3 = [
        '基于 PyTorch 构建融合 LSTM + 注意力机制的时空特征神经网络，适配经纬度区域映射数据集，完成数据预处理、模型训练、效果验证全流程；',
        '搭建训练监控体系，实现 MAE/RMSE 指标实时追踪、断点续训、最优权重保存与训练日志可视化，保障模型迭代效率；',
        '主导模型结构与超参数调优，通过正则化、早停机制解决过拟合问题，使模型预测精度较初始版本提升 15%+，达成团队既定指标；',
        '完成多场景测试验证，输出标准化推理流程与模型使用文档，高效交付可用模型，为团队业务落地提供核心 AI 能力支撑。'
    ]
    
    for resp in responsibilities3:
        doc.add_paragraph(resp, style='List Number')
    
    doc.add_paragraph()
    
    # 项目4：智能问诊系统项目（原有，修改为Agent版本）
    doc.add_heading('智能问诊系统项目', 2)
    p4 = doc.add_paragraph()
    p4.add_run('AI 算法 Agent（负责人）\n').bold = True
    
    responsibilities4 = [
        '基于 PyTorch + GPT2/BERT 框架构建对话生成模型，适配医疗问答语料数据集，完成数据预处理、模型训练、服务部署全流程搭建；',
        '搭建标准化数据处理流水线，通过分词编码、数据分割、批次封装实现医疗对话数据规范化，设计 collate_fn 函数完成序列填充与标签适配，保障数据输入一致性；',
        '主导模型优化与训练体系搭建：采用 AdamW 优化器、学习率预热与梯度裁剪策略，构建 "训练-验证" 双阶段流程，实时监控损失指标，通过正则化抑制过拟合，模型生成效果满足医疗对话交互需求；',
        '完成工程化部署与多场景验证：开发 Flask API 接口与本地交互程序，支持跨域请求与实时对话生成，输出标准化使用文档，为智能问诊业务落地提供核心 AI 交互能力支撑。'
    ]
    
    for resp in responsibilities4:
        doc.add_paragraph(resp, style='List Number')
    
    # 保存
    output_path = 'resume_updated.docx'
    doc.save(output_path)
    print(f'✅ 简历已生成: {output_path}')
    
    # 复制到桌面
    import shutil
    desktop_path = os.path.expanduser('~/Desktop/简历_已更新.docx')
    shutil.copy(output_path, desktop_path)
    print(f'✅ 已复制到桌面: {desktop_path}')

if __name__ == '__main__':
    import os
    create_resume()
